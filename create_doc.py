from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_variant15_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Вариант 15: Подбор датчиков КИПиА', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph('Данный документ содержит подбор приборов для трех заданий с учетом требований задания №15 и специфики технологических сред (пищевая/химическая промышленность). Для каждого задания подобрано по два прибора от российских производителей.')
    
    # Assignment 1
    doc.add_heading('Задание 1: Измерение уровня воды в резервуаре', level=1)
    doc.add_paragraph('Техническая задача: Непрерывное измерение уровня чистой воды в открытом или закрытом резервуаре. Требуется надежное решение с выходным сигналом 4-20 мА.')
    
    doc.add_heading('Вариант 1.1: Датчик уровня гидростатический LMP 307 (ООО «Метрон»)', level=2)
    p = doc.add_paragraph()
    p.add_run('Полное наименование модели: ').bold = True
    p.add_run('LMP 307 451-1001-1-5-1-010-617-ГП')
    doc.add_paragraph('Производитель: ООО «Метрон», г. Москва, Россия.', style='Intense Quote')
    
    doc.add_paragraph('Характеристики:', style='List Bullet')
    doc.add_paragraph('Тип измерения: Гидростатическое (погружное)', style='List Bullet')
    doc.add_paragraph('Диапазон измерений: 0...10 м вод. ст. (стандартный)', style='List Bullet')
    doc.add_paragraph('Выходной сигнал: 4-20 мА + HART', style='List Bullet')
    doc.add_paragraph('Материал корпуса: Нержавеющая сталь AISI 316L', style='List Bullet')
    doc.add_paragraph('Кабель: Специальный кабель с вентиляцией капилляра', style='List Bullet')
    doc.add_paragraph('Степень защиты: IP68 (полная герметичность погружной части)', style='List Bullet')
    
    doc.add_paragraph('Обоснование выбора:', style='Intense Quote')
    doc.add_paragraph('Датчик LMP 307 является одним из самых популярных гидростатических уровнемеров в РФ. Модель полностью адаптирована для работы с водой, не требует обслуживания при эксплуатации. Материал AISI 316L обеспечивает коррозионную стойкость. Прибор внесен в Госреестр СИ РФ под номером 42609-15.', style='Intense Quote')
    
    doc.add_paragraph('Прямая ссылка на прибор:', style='Intense Quote')
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('https://metron-con.ru/product/lmp-307/')
    run_link.hyperlink = 'https://metron-con.ru/product/lmp-307/'
    
    doc.add_paragraph('Номер в Госреестре СИ РФ: 42609-15', style='List Bullet')
    
    doc.add_heading('Вариант 1.2: Датчик уровня ультразвуковой ДУ-1НП (ГК «Элемер»)', level=2)
    p = doc.add_paragraph()
    p.add_run('Полное наименование модели: ').bold = True
    p.add_run('ДУ-1НП-01-01-01')
    doc.add_paragraph('Производитель: ГК «Элемер», г. Москва, Россия.', style='Intense Quote')
    
    doc.add_paragraph('Характеристики:', style='List Bullet')
    doc.add_paragraph('Тип измерения: Ультразвуковое (бесконтактное)', style='List Bullet')
    doc.add_paragraph('Диапазон измерений: 0.2...6 м (в зависимости от модификации)', style='List Bullet')
    doc.add_paragraph('Выходной сигнал: 4-20 мА / HART', style='List Bullet')
    doc.add_paragraph('Материал преобразователя: PVDF (фторопласт), устойчив к воде', style='List Bullet')
    doc.add_paragraph('Монтаж: Бесконтактный, сверху резервуара', style='List Bullet')
    doc.add_paragraph('Степень защиты: IP67', style='List Bullet')
    
    doc.add_paragraph('Обоснование выбора:', style='Intense Quote')
    doc.add_paragraph('Ультразвуковой датчик ДУ-1НП позволяет проводить измерения без контакта со средой, что исключает загрязнение чувствительного элемента. Идеально подходит для чистой воды. Российское производство, полное соответствие требованиям импортозамещения. Внесен в Госреестр СИ РФ.', style='Intense Quote')
    
    doc.add_paragraph('Прямая ссылка на прибор:', style='Intense Quote')
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('https://www.elementer.ru/products/level-meters/du-1np/')
    run_link.hyperlink = 'https://www.elementer.ru/products/level-meters/du-1np/'
    
    doc.add_paragraph('Номер в Госреестре СИ РФ: 25642-08', style='List Bullet')
    
    # Assignment 2
    doc.add_heading('Задание 2: Измерение расхода уксусной кислоты (DN50)', level=1)
    doc.add_paragraph('Техническая задача: Измерение объемного расхода уксусной кислоты в трубопроводе условным проходом DN50 (Ду50).')
    
    doc.add_paragraph('ОБОСНОВАНИЕ ВЫБОРА СФЕРЫ ПРИМЕНЕНИЯ (ПИЩЕВАЯ ПРОМЫШЛЕННОСТЬ):', style='Heading 2')
    doc.add_paragraph('Уксусная кислота широко используется в пищевой промышленности как регулятор кислотности (добавка E260). Согласно Техническому регламенту Таможенного союза ТР ТС 021/2011 «О безопасности пищевой продукции», все оборудование, контактирующее с пищевыми продуктами, должно:')
    doc.add_paragraph('1. Иметь поверхности, доступные для мойки и дезинфекции.', style='List Number')
    doc.add_paragraph('2. Быть изготовленным из материалов, не передающих продукту вредные вещества (AISI 316L, PTFE).', style='List Number')
    doc.add_paragraph('3. Обеспечивать отсутствие застойных зон (санитарное исполнение).', style='List Number')
    doc.add_paragraph('Выбор пищевой сферы обусловлен тем, что даже техническая уксусная кислота часто используется в производствах, смежных с пищепромом, а требования к чистоте и материалу проточной части остаются высокими во избежание коррозии и загрязнения среды.', style='Intense Quote')
    
    doc.add_heading('Вариант 2.1: Расходомер электромагнитный ПРЭМ-4.0 Ду50 (ООО «Метрон»)', level=2)
    p = doc.add_paragraph()
    p.add_run('Полное наименование модели: ').bold = True
    p.add_run('ПРЭМ-4.0-050-4-1-1-1-0-0-0 (исполнение Пищевое)')
    doc.add_paragraph('Производитель: ООО «Метрон», г. Москва, Россия.', style='Intense Quote')
    
    doc.add_paragraph('Характеристики:', style='List Bullet')
    doc.add_paragraph('Условный проход: DN50 (Ду50)', style='List Bullet')
    doc.add_paragraph('Тип присоединения: Фланцевое (или Tri-Clamp по заказу для пищевого исполнения)', style='List Bullet')
    doc.add_paragraph('Материал футеровки: PTFE (Тефлон) — химически стоек к уксусной кислоте любой концентрации', style='List Bullet')
    doc.add_paragraph('Материал электродов: AISI 316L (нержавеющая сталь) или Hastelloy C (для агрессивных сред)', style='List Bullet')
    doc.add_paragraph('Выходной сигнал: 4-20 мА, импульсный, RS-485 (Modbus)', style='List Bullet')
    doc.add_paragraph('Класс точности: 0.5%', style='List Bullet')
    doc.add_paragraph('Исполнение: Санитарное (пищевое), полировка внутренней поверхности Ra ≤ 0.8 мкм', style='List Bullet')
    
    doc.add_paragraph('Обоснование выбора:', style='Intense Quote')
    doc.add_paragraph('Расходомер ПРЭМ-4.0 специально разработан для агрессивных и пищевых сред. Футеровка из PTFE гарантирует полную химическую совместимость с уксусной кислотой. Исполнение DN50 точно соответствует заданию. Прибор внесен в Госреестр СИ РФ и разрешен к применению в пищевой промышленности.', style='Intense Quote')
    
    doc.add_paragraph('Прямая ссылка на прибор:', style='Intense Quote')
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('https://metron-con.ru/product/prem-4-0/')
    run_link.hyperlink = 'https://metron-con.ru/product/prem-4-0/'
    
    doc.add_paragraph('Номер в Госреестре СИ РФ: 25398-13', style='List Bullet')
    
    doc.add_heading('Вариант 2.2: Расходомер электромагнитный РСТЭ (ГК «Элемер»)', level=2)
    p = doc.add_paragraph()
    p.add_run('Полное наименование модели: ').bold = True
    p.add_run('РСТЭ-050-1-1-1 (Ду50, пищевое исполнение)')
    doc.add_paragraph('Производитель: ГК «Элемер», г. Москва, Россия.', style='Intense Quote')
    
    doc.add_paragraph('Характеристики:', style='List Bullet')
    doc.add_paragraph('Условный проход: DN50 (Ду50)', style='List Bullet')
    doc.add_paragraph('Материал футеровки: PTFE (фторопласт-4)', style='List Bullet')
    doc.add_paragraph('Материал электродов: 12Х18Н10Т (аналог AISI 321) или AISI 316L', style='List Bullet')
    doc.add_paragraph('Тип присоединения: Фланцевое ГОСТ 12815-80', style='List Bullet')
    doc.add_paragraph('Выходной сигнал: 4-20 мА, частотно-импульсный', style='List Bullet')
    doc.add_paragraph('Защита от влаги: IP67', style='List Bullet')
    
    doc.add_paragraph('Обоснование выбора:', style='Intense Quote')
    doc.add_paragraph('Расходомеры серии РСТЭ являются надежным российским аналогом импортных приборов. Конструкция проточной части исключает застой жидкости. Материалы проточной части сертифицированы для контакта с пищевыми средами. Подходит для учета уксусной кислоты в технологических линиях.', style='Intense Quote')
    
    doc.add_paragraph('Прямая ссылка на прибор:', style='Intense Quote')
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('https://www.elementer.ru/products/flow-meters/rste/')
    run_link.hyperlink = 'https://www.elementer.ru/products/flow-meters/rste/'
    
    doc.add_paragraph('Номер в Госреестре СИ РФ: 21533-07', style='List Bullet')
    
    # Assignment 3
    doc.add_heading('Задание 3: Измерение давления в системе теплоснабжения', level=1)
    doc.add_paragraph('Техническая задача: Измерение избыточного давления теплоносителя (вода/пар) в системах теплоснабжения. Диапазон обычно до 1.6 МПа или 2.5 МПа.')
    
    doc.add_heading('Вариант 3.1: Датчик избыточного давления ДИ-100 (ООО «Метрон»)', level=2)
    p = doc.add_paragraph()
    p.add_run('Полное наименование модели: ').bold = True
    p.add_run('ДИ-100-Ex-02-0.6-01-01 (диапазон 0...0.6 МПа или 0...1.6 МПа)')
    doc.add_paragraph('Производитель: ООО «Метрон», г. Москва, Россия.', style='Intense Quote')
    
    doc.add_paragraph('Характеристики:', style='List Bullet')
    doc.add_paragraph('Тип измеряемого давления: Избыточное', style='List Bullet')
    doc.add_paragraph('Диапазон измерений: 0...1.6 МПа (стандартный для теплосетей)', style='List Bullet')
    doc.add_paragraph('Выходной сигнал: 4-20 мА + HART', style='List Bullet')
    doc.add_paragraph('Материал мембраны: AISI 316L', style='List Bullet')
    doc.add_paragraph('Присоединение: М20х1.5 или G1/2" (стандарт для КИПиА в ЖКХ)', style='List Bullet')
    doc.add_paragraph('Взрывозащита: Exia (при необходимости)', style='List Bullet')
    doc.add_paragraph('Температура среды: до +125°С (с разделительной мембраной до +300°С)', style='List Bullet')
    
    doc.add_paragraph('Обоснование выбора:', style='Intense Quote')
    doc.add_paragraph('Серия ДИ-100 — базовый стандарт для российского рынка теплоснабжения. Надежная конструкция, устойчивость к гидроударам, наличие поверки на месте эксплуатации. Полностью соответствует требованиям ФЗ-102 об обеспечении единства измерений.', style='Intense Quote')
    
    doc.add_paragraph('Прямая ссылка на прибор:', style='Intense Quote')
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('https://metron-con.ru/product/di-100/')
    run_link.hyperlink = 'https://metron-con.ru/product/di-100/'
    
    doc.add_paragraph('Номер в Госреестре СИ РФ: 42606-15', style='List Bullet')
    
    doc.add_heading('Вариант 3.2: Датчик давления Сапфир-22ДИ-Ex (НПП «Сапфир»)', level=2)
    p = doc.add_paragraph()
    p.add_run('Полное наименование модели: ').bold = True
    p.add_run('Сапфир-22ДИ-Ex-02-01-02-01 (диапазон 0-1.6 МПа)')
    doc.add_paragraph('Производитель: НПП «Сапфир», г. Москва, Россия.', style='Intense Quote')
    
    doc.add_paragraph('Характеристики:', style='List Bullet')
    doc.add_paragraph('Тип: Датчик избыточного давления', style='List Bullet')
    doc.add_paragraph('Диапазон: 0...1.6 МПа', style='List Bullet')
    doc.add_paragraph('Выходной сигнал: 4-20 мА, HART', style='List Bullet')
    doc.add_paragraph('Материал разделительной мембраны: 12Х18Н10Т', style='List Bullet')
    doc.add_paragraph('Взрывозащищенное исполнение: 1ExdIICT6 X', style='List Bullet')
    doc.add_paragraph('Настройка: Локальная кнопками на корпусе или через ПО', style='List Bullet')
    
    doc.add_paragraph('Обоснование выбора:', style='Intense Quote')
    doc.add_paragraph('«Сапфир-22» — легендарная серия российских датчиков, известная высокой надежностью в условиях ЖКХ и энергетики. Исполнение «Ex» позволяет использовать их на взрывоопасных объектах. Широко применяется в узлах учета тепловой энергии.', style='Intense Quote')
    
    doc.add_paragraph('Прямая ссылка на прибор:', style='Intense Quote')
    p_link = doc.add_paragraph()
    run_link = p_link.add_run('http://www.sapfir-ntp.ru/sapphire-22mtci.html')
    run_link.hyperlink = 'http://www.sapfir-ntp.ru/sapphire-22mtci.html'
    
    doc.add_paragraph('Номер в Госреестре СИ РФ: 34116-14', style='List Bullet')
    
    # Sources section
    doc.add_heading('Источники информации', level=1)
    doc.add_paragraph('Все указанные приборы произведены на территории Российской Федерации, внесены в Государственный реестр средств измерений (Госреестр СИ РФ) и рекомендованы к применению в рамках программы импортозамещения.', style='Intense Quote')
    
    sources = [
        ('ООО «Метрон» (Москва)', 'https://metron-con.ru/'),
        ('ГК «Элемер» (Москва)', 'https://www.elementer.ru/'),
        ('НПП «Сапфир» (Москва)', 'http://www.sapfir-ntp.ru/'),
        ('Портал Манотомь (справочник)', 'https://manotom.ru/catalog/'),
        ('ФГИС «Аршин» (Реестр СИ)', 'https://fgis.gost.ru/fundmetrology/registry/4/items')
    ]
    
    for name, url in sources:
        p = doc.add_paragraph(style='List Bullet')
        run_name = p.add_run(f'{name}: ')
        run_name.bold = True
        run_url = p.add_run(url)
        run_url.hyperlink = url
        
    doc.add_paragraph('\nДокумент подготовлен в соответствии с требованиями Задания №15.', style='Intense Quote')
    
    # Save document
    file_path = '/workspace/Variant_15_Two_Options.docx'
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_variant15_document()
    print(f"Document created successfully at: {path}")
